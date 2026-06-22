from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from pdf2docx import Converter
from PIL import Image, ImageEnhance, ImageFilter

# Set U2NET_HOME to /tmp BEFORE importing rembg so the model
# downloads to a writable location on Render's filesystem
import os
os.environ.setdefault("U2NET_HOME", "/tmp/u2net")

from rembg import remove, new_session

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import fitz

import cv2
import numpy as np
import base64
import uuid
import io
import textwrap

app = FastAPI()

# Pre-initialize rembg session at startup so the model is downloaded
# once when the server boots, not on the first user request.
# This avoids a 30-60s delay (and potential timeout) on first use.
print("[startup] Loading rembg U2NET model...")
try:
    REMBG_SESSION = new_session("u2net")
    print("[startup] rembg model loaded successfully.")
except Exception as e:
    print(f"[startup] WARNING: rembg model failed to load: {e}")
    REMBG_SESSION = None


# -------------------------
# CORS
# -------------------------
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


# -------------------------
# HOME
# -------------------------
@app.get("/")
def home():
    return {
        "message": "ToolHubAI backend running 🚀",
        "status": "ok",
        "version": "1.0.0"
    }


# -------------------------
# HEALTH CHECK (for Render & uptime monitors)
# -------------------------
@app.get("/health")
def health():
    return {"status": "ok"}


# -------------------------
# OBJECT REMOVER
# -------------------------
@app.post("/remove-object")
async def remove_object(
    file: UploadFile = File(...),
    mask: UploadFile = File(...)
):
    image_bytes = await file.read()
    mask_bytes = await mask.read()

    np_image = np.frombuffer(image_bytes, np.uint8)
    np_mask = np.frombuffer(mask_bytes, np.uint8)

    image = cv2.imdecode(
        np_image,
        cv2.IMREAD_COLOR
    )

    mask_image = cv2.imdecode(
        np_mask,
        cv2.IMREAD_UNCHANGED
    )

    if image is None or mask_image is None:
        return {"error": "Invalid image or mask"}

    if len(mask_image.shape) == 3:
        mask_gray = cv2.cvtColor(
            mask_image,
            cv2.COLOR_BGR2GRAY
        )
    else:
        mask_gray = mask_image

    _, mask_thresh = cv2.threshold(
        mask_gray,
        10,
        255,
        cv2.THRESH_BINARY
    )

    inpainted = cv2.inpaint(
        image,
        mask_thresh,
        3,
        cv2.INPAINT_TELEA
    )

    _, buffer = cv2.imencode(
        ".png",
        inpainted
    )

    encoded = base64.b64encode(
        buffer
    ).decode("utf-8")

    return {"image": encoded}


# -------------------------
# CHAT ASSISTANT
# -------------------------
class ChatRequest(BaseModel):
    message: str


@app.post("/chat-assistant")
async def chat_assistant(request: ChatRequest):
    prompt = request.message.lower()

    if "deploy" in prompt:
        return {
            "answer": "Deploy frontend on Vercel and backend on Render. Add VITE_API_BASE_URL in your Vercel env settings pointing to your Render URL."
        }

    elif "pdf" in prompt or "convert" in prompt:
        return {
            "answer": "Use the PDF tools to convert files. Upload a PDF to get a DOCX, or upload a DOCX to get a PDF."
        }

    elif "background" in prompt or "remove bg" in prompt:
        return {
            "answer": "Use the Background Remover tool. Upload any image and the AI will strip the background and return a transparent PNG."
        }

    elif "object" in prompt or "erase" in prompt:
        return {
            "answer": "Use the Object Remover tool. Upload an image, paint over the object you want removed, then click Remove Object."
        }

    elif "enhance" in prompt or "sharpen" in prompt:
        return {
            "answer": "Use the AI Image Enhancer to boost sharpness, contrast, and clarity with one click."
        }

    elif "upscale" in prompt or "resolution" in prompt:
        return {
            "answer": "Use the Photo Upscaler to double the resolution of any image using AI interpolation."
        }

    elif "resume" in prompt or "cv" in prompt:
        return {
            "answer": "Use the Resume Analyzer to upload your resume PDF or DOCX and get a score, found skills, and ATS improvement tips."
        }

    return {
        "answer": "I can help you with ToolHubAI tools: Background Remover, Object Remover, Image Enhancer, Photo Upscaler, PDF Converter, and Resume Analyzer. Just ask me anything!"
    }


# -------------------------
# REMOVE BACKGROUND
# -------------------------
@app.post("/remove-bg")
async def remove_bg(
    file: UploadFile = File(...)
):
    try:
        contents = await file.read()

        if REMBG_SESSION is not None:
            # Use pre-loaded session (fast path)
            output = remove(contents, session=REMBG_SESSION)
        else:
            # Fallback: try loading model on-demand
            output = remove(contents)

        encoded = base64.b64encode(output).decode("utf-8")
        return {"image": encoded}

    except Exception as e:
        print(f"[remove-bg] Error: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Background removal failed: {str(e)}"
        )


# -------------------------
# PDF → DOCX
# -------------------------
@app.post("/pdf-to-docx")
async def pdf_to_docx(
    file: UploadFile = File(...)
):
    unique_id = str(uuid.uuid4())

    pdf_path = f"/tmp/{unique_id}.pdf"
    docx_path = f"/tmp/{unique_id}.docx"

    with open(pdf_path, "wb") as f:
        f.write(await file.read())

    converter = Converter(pdf_path)
    converter.convert(docx_path)
    converter.close()

    # Clean up source PDF
    if os.path.exists(pdf_path):
        os.remove(pdf_path)

    return FileResponse(
        path=docx_path,
        filename="converted.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=None
    )


# -------------------------
# DOCX → PDF
# -------------------------
@app.post("/docx-to-pdf")
async def docx_to_pdf(
    file: UploadFile = File(...)
):
    unique_id = str(uuid.uuid4())

    docx_path = f"/tmp/{unique_id}.docx"
    pdf_path = f"/tmp/{unique_id}.pdf"

    with open(docx_path, "wb") as f:
        f.write(await file.read())

    doc = Document(docx_path)
    pdf = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y = height - 72

    for para in doc.paragraphs:
        lines = textwrap.wrap(para.text, width=90) or [""]

        for line in lines:
            if y < 72:
                pdf.showPage()
                y = height - 72

            pdf.drawString(72, y, line)
            y -= 16

        y -= 8

    pdf.save()

    if os.path.exists(docx_path):
        os.remove(docx_path)

    return FileResponse(
        path=pdf_path,
        filename="converted.pdf",
        media_type="application/pdf"
    )


# -------------------------
# IMAGE ENHANCER
# -------------------------
@app.post("/enhance-image")
async def enhance_image(
    file: UploadFile = File(...)
):
    contents = await file.read()

    image = Image.open(
        io.BytesIO(contents)
    )

    image = image.filter(
        ImageFilter.SHARPEN
    )

    image = ImageEnhance.Contrast(
        image
    ).enhance(1.3)

    image = ImageEnhance.Sharpness(
        image
    ).enhance(2.0)

    buffer = io.BytesIO()

    image.save(
        buffer,
        format="PNG"
    )

    encoded = base64.b64encode(
        buffer.getvalue()
    ).decode("utf-8")

    return {"image": encoded}


# -------------------------
# RESUME ANALYZER
# -------------------------
@app.post("/analyze-resume")
async def analyze_resume(
    file: UploadFile = File(...)
):
    text = ""
    unique_id = str(uuid.uuid4())

    if file.filename.endswith(".pdf"):
        contents = await file.read()

        pdf_path = f"/tmp/resume_{unique_id}.pdf"
        with open(pdf_path, "wb") as f:
            f.write(contents)

        with fitz.open(pdf_path) as pdf:
            for page in pdf:
                extracted = page.get_text()

                if extracted:
                    text += extracted

        if os.path.exists(pdf_path):
            os.remove(pdf_path)

    elif file.filename.endswith(".docx"):
        contents = await file.read()

        docx_path = f"/tmp/resume_{unique_id}.docx"
        with open(docx_path, "wb") as f:
            f.write(contents)

        doc = Document(docx_path)

        for para in doc.paragraphs:
            text += para.text + "\n"

        if os.path.exists(docx_path):
            os.remove(docx_path)

    skills_db = [
        "python",
        "java",
        "sql",
        "react",
        "fastapi",
        "javascript",
        "html",
        "css",
        "git",
        "github",
        "typescript",
        "node",
        "docker",
        "aws",
        "flask",
        "django",
        "machine learning",
        "data science",
        "tensorflow",
        "pytorch",
    ]

    found_skills = []

    for skill in skills_db:
        if skill in text.lower():
            found_skills.append(skill)

    score = min(
        100,
        len(found_skills) * 7 + 20
    )

    return {
        "score": score,
        "skills": found_skills,
        "suggestions": [
            "Add a Projects section with links",
            "Use ATS-friendly keywords from job descriptions",
            "Quantify achievements with numbers (e.g., 'improved speed by 30%')",
            "Keep to 1 page if under 5 years experience"
        ],
        "ats_tips": [
            "Use a clear summary with your target role",
            "Add measurable outcomes to experience bullets",
            "Match important keywords from the job description",
            "Avoid tables, images, or columns — ATS can't parse them",
            "Save and submit as PDF unless otherwise requested"
        ]
    }


# -------------------------
# PHOTO UPSCALER
# -------------------------
@app.post("/upscale-image")
async def upscale_image(
    file: UploadFile = File(...)
):
    contents = await file.read()

    np_arr = np.frombuffer(
        contents,
        np.uint8
    )

    image = cv2.imdecode(
        np_arr,
        cv2.IMREAD_COLOR
    )

    upscaled = cv2.resize(
        image,
        None,
        fx=2,
        fy=2,
        interpolation=cv2.INTER_CUBIC
    )

    _, buffer = cv2.imencode(
        ".png",
        upscaled
    )

    encoded = base64.b64encode(
        buffer
    ).decode("utf-8")

    return {"image": encoded}
